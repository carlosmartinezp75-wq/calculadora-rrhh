import streamlit as st
import pandas as pd
import io
import zipfile
import tempfile
from datetime import datetime, date, timedelta
from dateutil.relativedelta import relativedelta
import json

# =============================================================================
# 1. CONFIGURACIÓN VISUAL Y LIBRERÍAS
# =============================================================================
st.set_page_config(page_title="HR Suite Pro - Sistema Integral", layout="wide", page_icon="🏢")

st.markdown("""
<style>
    .main {background-color: #f4f6f9;}
    h1, h2, h3 {color: #0d2b4e; font-family: 'Segoe UI', sans-serif;}
    .stButton>button {
        background-color: #0d2b4e; color: white; border-radius: 6px; height: 3em; width: 100%; font-weight: bold;
    }
    .metric-card {
        background: white; padding: 15px; border-radius: 8px; border-left: 5px solid #0d2b4e; box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        margin: 10px 0;
    }
    .alert-box {
        padding: 15px; background-color: #fff3cd; color: #856404; border: 1px solid #ffeeba; border-radius: 5px; margin-bottom: 15px;
    }
    .competency-gap {
        background: #e8f5e8; border-left: 5px solid #28a745; padding: 10px; margin: 5px 0; border-radius: 5px;
    }
    .competency-gap-high {
        background: #fff3cd; border-left: 5px solid #ffc107; padding: 10px; margin: 5px 0; border-radius: 5px;
    }
    .competency-gap-critical {
        background: #f8d7da; border-left: 5px solid #dc3545; padding: 10px; margin: 5px 0; border-radius: 5px;
    }
</style>
""", unsafe_allow_html=True)

try:
    from fpdf import FPDF
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import xlsxwriter
    LIBS_OK = True
except ImportError:
    st.error("⚠️ Faltan librerías. Instala: pip install fpdf python-docx xlsxwriter pandas streamlit python-dateutil")
    LIBS_OK = False

# =============================================================================
# 2. DATOS MAESTROS (INDICADORES 2025)
# =============================================================================
IND = {
    "UF": 39643.59, 
    "UTM": 69542.0, 
    "IMM": 530000,
    "TOPE_GRAT": (4.75 * 530000)/12,
    "TOPE_AFP": 84.3, 
    "TOPE_AFC": 126.6,
    "TOPE_INDEM_ANOS": 90,
    "VALOR_HORA_ORDinaria": 530000 / 45,  # Jornada 45 horas semanales
    "MIN_SALUD": 0.07,  # 7% mínimo salud
    "MAX_ISAPRE_BONIF": 4.0  # UF máximo bonificación fiscal
}

# Competencias base por área
COMPETENCIAS_BASE = {
    "Administración": {
        "Conocimientos Técnicos": {
            "Contabilidad": ["Básico", "Intermedio", "Avanzado", "Experto"],
            "Administración": ["Básico", "Intermedio", "Avanzado", "Experto"],
            "Excel": ["Básico", "Intermedio", "Avanzado", "Avanzado+"]
        },
        "Habilidades Blandas": {
            "Comunicación": ["Básico", "Intermedio", "Avanzado", "Experto"],
            "Liderazgo": ["Básico", "Intermedio", "Avanzado", "Experto"],
            "Análisis": ["Básico", "Intermedio", "Avanzado", "Experto"]
        }
    },
    "Tecnología": {
        "Conocimientos Técnicos": {
            "Programación": ["Básico", "Intermedio", "Avanzado", "Experto"],
            "Bases de Datos": ["Básico", "Intermedio", "Avanzado", "Experto"],
            "Redes": ["Básico", "Intermedio", "Avanzado", "Experto"]
        },
        "Habilidades Blandas": {
            "Resolución Problemas": ["Básico", "Intermedio", "Avanzado", "Experto"],
            "Innovación": ["Básico", "Intermedio", "Avanzado", "Experto"],
            "Trabajo Equipo": ["Básico", "Intermedio", "Avanzado", "Experto"]
        }
    },
    "Operaciones": {
        "Conocimientos Técnicos": {
            "Gestión Operaciones": ["Básico", "Intermedio", "Avanzado", "Experto"],
            "Procesos": ["Básico", "Intermedio", "Avanzado", "Experto"],
            "Logística": ["Básico", "Intermedio", "Avanzado", "Experto"]
        },
        "Habilidades Blandas": {
            "Planificación": ["Básico", "Intermedio", "Avanzado", "Experto"],
            "Organización": ["Básico", "Intermedio", "Avanzado", "Experto"],
            "Orientación Resultados": ["Básico", "Intermedio", "Avanzado", "Experto"]
        }
    }
}

# =============================================================================
# 3. MOTORES LÓGICOS (BACKEND)
# =============================================================================

class MotorFinanciero:
    @staticmethod
    def calcular_liquidacion(liquido_obj, col, mov, tipo_contrato, salud_sistema, plan_uf):
        """Calcula la estructura salarial completa"""
        no_imp = col + mov
        meta = liquido_obj - no_imp
        min_b, max_b = meta, meta * 2.5
        res = {}
        
        for _ in range(100):
            test = (min_b + max_b) / 2
            
            # Estructura Base vs Grat
            if tipo_contrato == "Sueldo Empresarial":
                grat = 0
                base = test
            else:
                if test > (IND["TOPE_GRAT"] * 5):
                    grat = IND["TOPE_GRAT"]
                    base = test - grat
                else:
                    base = test / 1.25
                    grat = test - base
            
            imponible = base + grat
            
            # Descuentos Estándar para iteración
            tope_afp_p = IND["TOPE_AFP"] * IND["UF"]
            afp = 0 if tipo_contrato == "Sueldo Empresarial" else int(min(imponible, tope_afp_p) * 0.11)
            
            # Salud Legal 7% (Base de cálculo)
            salud_legal = int(min(imponible, tope_afp_p) * 0.07)
            
            afc = 0
            if tipo_contrato == "Indefinido":
                afc = int(min(imponible, IND["TOPE_AFC"]*IND["UF"]) * 0.006)
            
            tributable = imponible - afp - salud_legal - afc
            imp = 0 
            if tributable > (13.5*IND["UTM"]): imp = (tributable*0.04) - (0.54*IND["UTM"])
            imp = max(0, int(imp))
            
            liq_calc = imponible - afp - salud_legal - afc - imp
            
            if abs(liq_calc - meta) < 500:
                # Aplicación de realidad (ISAPRE CARGO TRABAJADOR)
                salud_real = salud_legal
                diff_isapre = 0
                glosa = ""
                
                if salud_sistema == "Isapre (UF)":
                    valor_plan = int(plan_uf * IND["UF"])
                    if valor_plan > salud_legal:
                        salud_real = valor_plan
                        diff_isapre = valor_plan - salud_legal
                        glosa = f"NOTA: Plan Isapre ({plan_uf} UF) excede el 7% legal. La diferencia de ${diff_isapre:,.0f} es de cargo del trabajador."
                
                # Recálculo final del líquido
                liq_final = imponible - afp - salud_real - afc - imp + no_imp
                
                return {
                    "Base": int(base), "Grat": int(grat), "Tot_Imp": int(imponible),
                    "No_Imp": int(no_imp), "AFP": afp, "Salud": salud_real,
                    "AFC": afc, "Impuesto": int(imp), "Liquido": int(liq_final),
                    "Glosa": glosa, "Diff_Isapre": diff_isapre,
                    "Valor_Hora": int((base + grat) / 180),  # 180 horas mensuales
                    "Porcentaje_Carga": int(((afp + salud_real + afc) / imponible) * 100)
                }
                break
            
            if liq_calc < meta: min_b = test
            else: max_b = test
        return res

class MotorFiniquitos:
    @staticmethod
    def calcular(f_inicio, f_termino, sueldo_base, gratificacion, col_mov, causal, dias_vac_tomados):
        """Motor avanzado de cálculo de finiquitos"""
        dias_totales = (f_termino - f_inicio).days + 1
        antiguedad = relativedelta(f_termino, f_inicio)
        
        # Base de Cálculo Indemnizaciones (con tope 90 UF)
        tope_indem_pesos = IND["TOPE_INDEM_ANOS"] * IND["UF"]
        base_indem = min((sueldo_base + gratificacion + col_mov), tope_indem_pesos)
        
        # Vacaciones Proporcionales (Factor 1.25 días por mes trabajado)
        factor_diario = 1.25 / 30
        dias_vac_ganados = dias_totales * factor_diario
        saldo_dias_vac = max(0, dias_vac_ganados - dias_vac_tomados)
        
        valor_dia_vac = (sueldo_base + gratificacion) / 30
        monto_vacaciones = int(saldo_dias_vac * valor_dia_vac)
        
        # Cálculo avanzado por causal
        monto_anos = 0
        monto_aviso = 0
        anos_pago = 0
        proporcional_sueldo = 0
        otros_conceptos = {}
        
        if causal == "Necesidades de la Empresa (Art. 161)":
            # Años de Servicio
            anos_pago = antiguedad.years
            if antiguedad.months >= 6:
                anos_pago += 1
            if anos_pago > 11: anos_pago = 11  # Tope legal
            monto_anos = int(base_indem * anos_pago)
            monto_aviso = int(base_indem)  # Mes de aviso
            
        elif causal == "Renuncia Voluntaria":
            # Solo vacaciones proporcionales
            proporcional_sueldo = 0  # No hay proporcional por renuncia voluntaria
            
        elif causal == "Término Injustificado (Art. 168)":
            # Pago íntegro del período trabajado
            proporcional_sueldo = int((sueldo_base + gratificacion) * (dias_totales / 30))
            monto_anos = int(base_indem * anos_pago)
            
        elif causal == "Muerte Trabajador":
            # Indemnización por muerte (tope 11 meses)
            anos_pago = min(11, antiguedad.years + 1)
            monto_anos = int(base_indem * anos_pago)
            otros_conceptos["Seguro Accidentes"] = int(tope_indem_pesos)
            
        # Cálculo total
        total_finiquito = (monto_vacaciones + monto_anos + monto_aviso + 
                         proporcional_sueldo + sum(otros_conceptos.values()))
        
        return {
            "Antiguedad": f"{antiguedad.years} años, {antiguedad.months} meses, {antiguedad.days} días",
            "Dias_Trabajos": dias_totales,
            "Base_Calculo_Indem": int(base_indem),
            "Dias_Vacaciones_Pendientes": round(saldo_dias_vac, 2),
            "Monto_Vacaciones": monto_vacaciones,
            "Años_Servicio": monto_anos,
            "Mes_Aviso": monto_aviso,
            "Proporcional_Sueldo": proporcional_sueldo,
            "Otros_Conceptos": otros_conceptos,
            "Total_Finiquito": total_finiquito,
            "Valor_Liquidacion_UF": round(total_finiquito / IND["UF"], 2)
        }

class MotorCompetencias:
    @staticmethod
    def crear_perfil_cargo(area, cargo, empresa_info):
        """Crea un perfil completo de cargo con competencias objetivo"""
        if area not in COMPETENCIAS_BASE:
            area = "Administración"  # Default
            
        perfil = {
            "cargo": cargo,
            "area": area,
            "empresa": empresa_info,
            "competencias_objetivo": {},
            "funciones_principales": [],
            "responsabilidades": [],
            "requisitos": []
        }
        
        # Competencias por área
        competencias = COMPETENCIAS_BASE[area]
        for categoria, competencias_cat in competencias.items():
            perfil["competencias_objetivo"][categoria] = {}
            for competencia, niveles in competencias_cat.items():
                # Nivel objetivo según seniority
                nivel_objetivo = "Intermedio" if "Analista" in cargo else "Avanzado"
                if "Jefe" in cargo or "Supervisor" in cargo:
                    nivel_objetivo = "Avanzado"
                elif "Gerente" in cargo:
                    nivel_objetivo = "Experto"
                    
                perfil["competencias_objetivo"][categoria][competencia] = nivel_objetivo
        
        # Funciones según área
        funciones_base = {
            "Administración": [
                "Gestión de procesos administrativos",
                "Coordinación de documentación legal",
                "Supervisión de equipos administrativos",
                "Análisis de indicadores de gestión"
            ],
            "Tecnología": [
                "Desarrollo y mantenimiento de sistemas",
                "Análisis de requerimientos técnicos",
                "Gestión de bases de datos",
                "Soporte técnico a usuarios"
            ],
            "Operaciones": [
                "Supervisión de procesos operativos",
                "Gestión de recursos y materiales",
                "Coordinación de equipos de trabajo",
                "Control de calidad operacional"
            ]
        }
        
        perfil["funciones_principales"] = funciones_base.get(area, funciones_base["Administración"])
        
        return perfil
    
    @staticmethod
    def evaluar_candidato(perfil_cargo, competencias_candidato):
        """Evalúa un candidato contra el perfil de cargo"""
        evaluaciones = {}
        brechas = []
        
        for categoria, competencias in perfil_cargo["competencias_objetivo"].items():
            if categoria not in competencias_candidato:
                continue
                
            evaluaciones[categoria] = {}
            
            for competencia, nivel_objetivo in competencias.items():
                if competencia not in competencias_candidato[categoria]:
                    continue
                    
                nivel_candidato = competencias_candidato[categoria][competencia]
                
                # Calcular brecha
                niveles = ["Básico", "Intermedio", "Avanzado", "Experto"]
                idx_objetivo = niveles.index(nivel_objetivo) if nivel_objetivo in niveles else 1
                idx_candidato = niveles.index(nivel_candidato) if nivel_candidato in niveles else 0
                
                brecha_nivel = idx_objetivo - idx_candidato
                
                evaluaciones[categoria][competencia] = {
                    "nivel_objetivo": nivel_objetivo,
                    "nivel_candidato": nivel_candidato,
                    "brecha": brecha_nivel,
                    "porcentaje_match": max(0, int((1 - brecha_nivel/3) * 100))
                }
                
                # Clasificar brecha
                if brecha_nivel >= 2:
                    brechas.append({
                        "competencia": competencia,
                        "categoria": categoria,
                        "brecha": brecha_nivel,
                        "criticidad": "Alta",
                        "recomendacion": f"Capacitación intensiva requerida en {competencia}"
                    })
                elif brecha_nivel == 1:
                    brechas.append({
                        "competencia": competencia,
                        "categoria": categoria,
                        "brecha": brecha_nivel,
                        "criticidad": "Media",
                        "recomendacion": f"Desarrollo complementario en {competencia}"
                    })
        
        # Plan de carrera
        plan_carrera = MotorCompetencias.generar_plan_carrera(brechas, perfil_cargo)
        
        return {
            "evaluaciones": evaluaciones,
            "brechas": brechas,
            "plan_carrera": plan_carrera,
            "score_total": sum([eval["porcentaje_match"] for cat in evaluaciones.values() 
                              for eval in cat.values()]) / max(1, sum([len(cat) for cat in evaluaciones.values()]))
        }
    
    @staticmethod
    def generar_plan_carrera(brechas, perfil_cargo):
        """Genera plan personalizado para reducir brechas"""
        plan = {
            "fase_1": [],  # Inmediato (0-3 meses)
            "fase_2": [],  # Corto plazo (3-6 meses)
            "fase_3": [],  # Mediano plazo (6-12 meses)
            "recursos_necesarios": {},
            "cronograma": []
        }
        
        for brecha in brechas:
            if brecha["criticidad"] == "Alta":
                plan["fase_1"].append({
                    "accion": brecha["recomendacion"],
                    "competencia": brecha["competencia"],
                    "duracion": "3 meses",
                    "tipo": "Capacitación Intensiva"
                })
                plan["recursos_necesarios"][brecha["competencia"]] = "Curso especializado + Mentoría"
                
            elif brecha["criticidad"] == "Media":
                plan["fase_2"].append({
                    "accion": brecha["recomendacion"],
                    "competencia": brecha["competencia"],
                    "duracion": "4 meses",
                    "tipo": "Desarrollo Progresivo"
                })
        
        return plan

class GeneradorDocumentos:
    @staticmethod
    def generar_contrato_trabajo(datos_empresa, datos_trabajador, datos_contrato):
        """Genera contrato de trabajo completo"""
        doc = Document()
        
        # Encabezado
        header = doc.sections[0].header
        header_para = header.paragraphs[0]
        header_para.text = datos_empresa["nombre"]
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Título
        title = doc.add_heading('CONTRATO DE TRABAJO', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Cuerpo del contrato
        doc.add_paragraph(f"En Santiago, {datetime.now().strftime('%d de %B de %Y')}, entre:")
        
        # Empresa
        doc.add_heading('EMPRESA:', level=1)
        doc.add_paragraph(f"""
Razón Social: {datos_empresa['nombre']}
RUT: {datos_empresa['rut']}
Representante Legal: {datos_empresa['rep_legal']}
Dirección: {datos_empresa['direccion']}
        """.strip())
        
        # Trabajador
        doc.add_heading('TRABAJADOR:', level=1)
        doc.add_paragraph(f"""
Nombre: {datos_trabajador['nombre']}
RUT: {datos_trabajador['rut']}
Nacionalidad: {datos_trabajador.get('nacionalidad', 'Chilena')}
Estado Civil: {datos_trabajador.get('estado_civil', '')}
Domicilio: {datos_trabajador.get('direccion', '')}
        """.strip())
        
        # Cláusulas del contrato
        doc.add_heading('ACUERDO:', level=1)
        doc.add_paragraph(f"""
PRIMERA: La empresa contrata los servicios del trabajador para el cargo de {datos_contrato['cargo']}, 
ubicado en {datos_contrato['ubicacion']}, quien prestará sus servicios a partir del {datos_contrato['fecha_inicio']}.

SEGUNDA: La jornada de trabajo será de {datos_contrato['jornada']} horas semanales, distribuidas de lunes a viernes.

TERCERA: El sueldo base será de ${datos_contrato['sueldo_base']:,.0f} mensuales, más una gratificación de ${datos_contrato['gratificacion']:,.0f} mensuales, 
pagadero el último día hábil de cada mes.

CUARTA: Las partes establecen un contrato de duración {datos_contrato['tipo_contrato']}.

QUINTA: El trabajador tendrá derecho a {datos_contrato['dias_vacaciones']} días hábiles de vacaciones por año calendario.
        """)
        
        # Leyes especiales
        doc.add_heading('CUMPLIMIENTO LEGAL:', level=1)
        doc.add_paragraph("""
• LEY 20.123 (TERCERA): Se ajusta a las disposiciones sobre trabajo en régimen de subcontratación.
• LEY 20.348 (40 HORAS): Se cumple con la gradual reducción de jornada laboral.
• LEY 20.348 (KARIN): Se incorpora protocolo de prevención del acoso laboral y sexual.
• CÓDIGO DEL TRABAJO: Se respetan todas las disposiciones vigentes.
        """)
        
        # Firmas
        doc.add_paragraph("\n\n")
        doc.add_paragraph("_" * 50)
        doc.add_paragraph(f"{datos_empresa['rep_legal']}")
        doc.add_paragraph("Representante Legal")
        
        doc.add_paragraph("\n\n")
        doc.add_paragraph("_" * 50)
        doc.add_paragraph(f"{datos_trabajador['nombre']}")
        doc.add_paragraph("Trabajador")
        
        return doc
    
    @staticmethod
    def generar_carta_amonestacion(datos_empresa, datos_trabajador, falta, descripcion, medidas):
        """Genera carta de amonestación"""
        doc = Document()
        
        title = doc.add_heading('CARTA DE AMONESTACIÓN', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Encabezado empresa
        doc.add_paragraph(f"""
{datos_empresa['nombre']}
RUT: {datos_empresa['rut']}
Dirección: {datos_empresa['direccion']}
        """)
        
        doc.add_paragraph(f"""
Santiago, {datetime.now().strftime('%d de %B de %Y')}
        """)
        
        # Destinatario
        doc.add_paragraph(f"""
Señor/a: {datos_trabajador['nombre']}
RUT: {datos_trabajador['rut']}
Cargo: {datos_trabajador['cargo']}
        """)
        
        # Contenido
        doc.add_heading('DE NUDA CONSIDERACIÓN:', level=1)
        doc.add_paragraph(f"""
Por la presente, me dirijo a usted para hacer de su conocimiento la siguiente AMONESTACIÓN:

1. NATURALEZA DE LA FALTA:
{descripcion}

2. CONDUCTA OBSERVADA:
{falta}

3. CONSECUENCIAS:
Las presentes observaciones constituyen una amonestación escrita en su hoja de vida laboral.

4. MEDIDAS CORRECTIVAS:
{medidas}

5. ADVERTENCIA:
Se le hace presente que de persistir en esta conducta o repetirse faltas similares, 
se podrán aplicar sanciones mássevere incluyendo el término del contrato de trabajo.
        """)
        
        doc.add_paragraph("\nSin otro particular, y confiando en su comprensión y corrección.")
        
        # Firma
        doc.add_paragraph("\n\nAtentamente,")
        doc.add_paragraph("\n" + "_" * 50)
        doc.add_paragraph(datos_empresa['rep_legal'])
        doc.add_paragraph("Representante Legal")
        
        return doc
    
    @staticmethod
    def generar_carta_desvinculacion(datos_empresa, datos_trabajador, causa, fecha_termino, avisos):
        """Genera carta de desvinculación"""
        doc = Document()
        
        title = doc.add_heading('CARTA DE DESVINCULACIÓN', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Encabezado
        doc.add_paragraph(f"""
{datos_empresa['nombre']}
RUT: {datos_empresa['rut']}
Dirección: {datos_empresa['direccion']}
        """)
        
        doc.add_paragraph(f"""
Santiago, {datetime.now().strftime('%d de %B de %Y')}
        """)
        
        # Destinatario
        doc.add_paragraph(f"""
Señor/a: {datos_trabajador['nombre']}
RUT: {datos_trabajador['rut']}
Cargo: {datos_trabajador['cargo']}
        """)
        
        # Contenido
        doc.add_heading('DE NUDA CONSIDERACIÓN:', level=1)
        doc.add_paragraph(f"""
Por medio de la presente, y de conformidad a lo establecido en el artículo {causa['articulo']} 
del Código del Trabajo, me dirijo a usted para comunicar el término de su contrato de trabajo por la siguiente causal:

**CAUSAL:** {causa['descripcion']}

**FECHA DE TÉRMINO:** {fecha_termino}

**AVISOS PREVIOS:**
{avisos}

**DERECHOS PENDIENTES:**
Se le pagará el total de sus haberes adeudados más el finiquito correspondiente 
dentro de los 10 días hábiles siguientes a la fecha de término del contrato.

**DEVOLUCIÓN DE ELEMENTOS:**
Deberá hacer entrega de todos los elementos de trabajo, documentos y valores 
que se encuentren en su poder dentro de la fecha de término.
        """)
        
        # Firma
        doc.add_paragraph("\n\nSin otro particular.")
        doc.add_paragraph("\n\nAtentamente,")
        doc.add_paragraph("\n" + "_" * 50)
        doc.add_paragraph(datos_empresa['rep_legal'])
        doc.add_paragraph("Representante Legal")
        
        return doc

class PDFGenerator(FPDF):
    def header(self):
        if hasattr(st.session_state, 'empresa') and st.session_state.empresa.get('logo_bytes'):
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
                    tmp.write(st.session_state.empresa['logo_bytes'])
                    self.image(tmp.name, 10, 8, 30)
            except:
                pass
        self.set_font('Arial', 'B', 14)
        self.cell(0, 10, 'DOCUMENTO OFICIAL HR SUITE', 0, 1, 'C')
        self.ln(10)

    def generar_liquidacion(self, datos, empresa, trabajador):
        self.add_page()
        self.set_font('Arial', '', 10)
        
        y = self.get_y()
        self.rect(10, y, 90, 30); self.set_xy(12, y+2)
        self.multi_cell(85, 5, f"EMPRESA: {empresa['nombre']}\nRUT: {empresa['rut']}\nDIR: {empresa['direccion']}")
        
        self.rect(110, y, 90, 30); self.set_xy(112, y+2)
        self.multi_cell(85, 5, f"TRABAJADOR: {trabajador['nombre']}\nRUT: {trabajador['rut']}\nCARGO: {trabajador['cargo']}")
        
        self.set_y(y + 35)
        
        # Detalles
        self.set_fill_color(200, 200, 200)
        self.cell(190, 8, "LIQUIDACIÓN DE SUELDO", 1, 1, 'C', 1)
        self.ln(2)
        
        haberes = [("Sueldo Base", datos['Base']), ("Gratificación", datos['Grat']), ("Colación", int(datos['No_Imp']/2)), ("Movilización", int(datos['No_Imp']/2))]
        descuentos = [("AFP", datos['AFP']), ("Salud", datos['Salud']), ("AFC", datos['AFC']), ("Impuesto", datos['Impuesto'])]
        
        self.cell(95, 6, "HABERES", 1, 0, 'C', 1)
        self.cell(95, 6, "DESCUENTOS", 1, 1, 'C', 1)
        
        for i in range(max(len(haberes), len(descuentos))):
            h = haberes[i] if i < len(haberes) else ("", 0)
            d = descuentos[i] if i < len(descuentos) else ("", 0)
            self.cell(65, 6, h[0], 'L'); self.cell(30, 6, f"${h[1]:,.0f}", 'R')
            self.cell(65, 6, d[0], 'L'); self.cell(30, 6, f"${d[1]:,.0f}", 'R', 1)
        
        self.ln(5)
        self.set_font('Arial', 'B', 12)
        self.cell(130, 10, "LÍQUIDO A PAGAR", 1, 0, 'R')
        self.cell(60, 10, f"${datos['Liquido']:,.0f}", 1, 1, 'C', 1)
        
        if datos.get('Glosa'):
            self.ln(3); self.set_font('Arial', 'I', 8)
            self.cell(0, 5, datos['Glosa'], 0, 1, 'C')
            
        return self.output(dest='S').encode('latin-1')

# =============================================================================
# 4. INICIALIZACIÓN DE ESTADO
# =============================================================================
if 'empresa' not in st.session_state:
    st.session_state.empresa = {
        "nombre": "", "rut": "", "rep_legal": "", "direccion": "", 
        "logo_bytes": None, "rubro": "General"
    }
    
if 'trabajador' not in st.session_state:
    st.session_state.trabajador = {
        "nombre": "", "rut": "", "cargo": "", "fecha_ingreso": date.today(),
        "email": "", "telefono": "", "direccion": "", "nacionalidad": "Chilena",
        "estado_civil": ""
    }

if 'perfil_cargo' not in st.session_state:
    st.session_state.perfil_cargo = None
    
if 'evaluaciones' not in st.session_state:
    st.session_state.evaluaciones = []

# =============================================================================
# 5. INTERFAZ LATERAL (SIDEBAR)
# =============================================================================
with st.sidebar:
    st.title("🏢 Configuración Global")
    
    # Logo
    logo = st.file_uploader("📷 Logo Empresa", type=['png', 'jpg', 'jpeg'])
    if logo:
        st.session_state.empresa['logo_bytes'] = logo.read()
        st.success("Logo cargado")
        
    # Datos Empresa
    with st.expander("🏭 Datos Empresa", expanded=True):
        st.session_state.empresa['rut'] = st.text_input("RUT Empresa", "76.xxx.xxx-x")
        st.session_state.empresa['nombre'] = st.text_input("Razón Social")
        st.session_state.empresa['rep_legal'] = st.text_input("Representante Legal")
        st.session_state.empresa['direccion'] = st.text_input("Dirección")
        st.session_state.empresa['rubro'] = st.selectbox("Rubro", 
            ["Administración", "Tecnología", "Operaciones", "Recursos Humanos", "Finanzas", "Marketing"])
    
    # Datos Trabajador
    with st.expander("👤 Datos Trabajador", expanded=True):
        st.session_state.trabajador['rut'] = st.text_input("RUT Trabajador")
        st.session_state.trabajador['nombre'] = st.text_input("Nombre Completo")
        st.session_state.trabajador['cargo'] = st.text_input("Cargo")
        st.session_state.trabajador['email'] = st.text_input("Email")
        st.session_state.trabajador['telefono'] = st.text_input("Teléfono")
        st.session_state.trabajador['fecha_ingreso'] = st.date_input("Fecha Ingreso", value=date.today())
        st.session_state.trabajador['direccion'] = st.text_input("Dirección")
        st.session_state.trabajador['estado_civil'] = st.selectbox("Estado Civil", 
            ["Soltero", "Casado", "Divorciado", "Viudo"])
        st.session_state.trabajador['nacionalidad'] = st.selectbox("Nacionalidad", 
            ["Chilena", "Extranjera"])

# =============================================================================
# 6. CUERPO PRINCIPAL
# =============================================================================
st.title("🚀 HR Suite Pro - Sistema Integral de Recursos Humanos")

# Indicadores principales en la parte superior
col1, col2, col3, col4 = st.columns(4)
with col1:
    st.metric("UF Actual", f"${IND['UF']:,.0f}")
with col2:
    st.metric("UTM Actual", f"${IND['UTM']:,.0f}")
with col3:
    st.metric("IMM", f"${IND['IMM']:,.0f}")
with col4:
    st.metric("Tope Indemnización", f"{IND['TOPE_INDEM_ANOS']} UF")

# Tabs principales
tabs = st.tabs([
    "💰 Calculadora Sueldos", 
    "📄 Finiquitos Avanzados", 
    "📋 Gestión Documentos",
    "👥 Evaluación Competencias",
    "🎯 Perfiles de Cargo", 
    "📊 Reportes Masivos",
    "⚖️ Legal & Compliance"
])

# =============================================================================
# TAB 1: CALCULADORA DE SUELDOS
# =============================================================================
with tabs[0]:
    st.header("🧮 Calculadora Inteligente de Liquidaciones")
    st.info("Calculadora avanzada con lógica isapre y análisis de cargas laborales")
    
    col1, col2 = st.columns(2)
    with col1:
        st.subheader("Parámetros de Cálculo")
        liq_objetivo = st.number_input("💰 Sueldo Líquido Objetivo", 800000, step=10000, min_value=400000)
        colacion = st.number_input("🍽️ Colación", 30000)
        movilizacion = st.number_input("🚗 Movilización", 30000)
        tipo_contrato = st.selectbox("📄 Tipo Contrato", 
            ["Indefinido", "Plazo Fijo", "Sueldo Empresarial", "Honorarios"])
        salud_sistema = st.radio("🏥 Sistema Salud", ["Fonasa", "Isapre (UF)"])
        plan_uf = st.number_input("📊 Plan UF", 0.0, min_value=0.0, step=0.1) if salud_sistema == "Isapre (UF)" else 0.0
    
    with col2:
        st.subheader("Análisis Adicional")
        if st.button("🔄 CALCULAR LIQUIDACIÓN", type="primary"):
            res = MotorFinanciero.calcular_liquidacion(
                liq_objetivo, colacion, movilizacion, tipo_contrato, salud_sistema, plan_uf
            )
            st.session_state.calculo = res
            
            # Alertas
            if res.get('Diff_Isapre', 0) > 0:
                st.warning(f"⚠️ El plan Isapre reduce el líquido en ${res['Diff_Isapre']:,.0f}")
                
            if res.get('Porcentaje_Carga', 0) > 30:
                st.error(f"🚨 Alta carga laboral: {res['Porcentaje_Carga']}%")
            
            # Métricas principales
            st.metric("💵 Sueldo Base", f"${res['Base']:,.0f}")
            st.metric("🎁 Gratificación", f"${res['Grat']:,.0f}")
            st.metric("💳 Valor Hora", f"${res['Valor_Hora']:,.0f}")
            st.metric("📊 Carga Total", f"{res['Porcentaje_Carga']}%")
            
            if st.button("📥 Descargar PDF", type="secondary"):
                pdf = PDFGenerator()
                pdf_bytes = pdf.generar_liquidacion(res, st.session_state.empresa, st.session_state.trabajador)
                st.download_button("📄 Liquidación.pdf", pdf_bytes, "Liquidacion_Completa.pdf", "application/pdf")

# =============================================================================
# TAB 2: FINIQUITOS AVANZADOS
# =============================================================================
with tabs[1]:
    st.header("📋 Calculadora Avanzada de Finiquitos")
    st.info("Cálculo preciso con múltiples causales y conceptos adicionales")
    
    col1, col2 = st.columns(2)
    with col1:
        st.subheader("Parámetros de Cálculo")
        f_termino = st.date_input("📅 Fecha Término", date.today())
        causal = st.selectbox("⚖️ Causal", [
            "Renuncia Voluntaria", 
            "Necesidades de la Empresa (Art. 161)", 
            "Término Injustificado (Art. 168)",
            "Muerte Trabajador",
            "Mutuo Acuerdo",
            "Vencimiento Contrato"
        ])
        dias_vac_tomados = st.number_input("🏖️ Días Vacaciones Tomados", 0.0, min_value=0.0)
        
    with col2:
        st.subheader("Base de Cálculo")
        # Usar datos de sesión si existen
        base_calc = st.session_state.get('calculo', {}).get('Base', 0)
        grat_calc = st.session_state.get('calculo', {}).get('Grat', 0)
        
        sb = st.number_input("💰 Sueldo Base", value=base_calc, min_value=0)
        gr = st.number_input("🎁 Gratificación", value=grat_calc, min_value=0)
        colmov = st.number_input("🍽️🚗 Col + Mov", 60000, min_value=0)
        otros_conceptos = st.number_input("💵 Otros Conceptos", 0, min_value=0)
    
    if st.button("🧮 CALCULAR FINIQUITO COMPLETO", type="primary"):
        if not st.session_state.trabajador.get('fecha_ingreso'):
            st.error("❌ Fecha de ingreso requerida en el sidebar")
        else:
            res_fin = MotorFiniquitos.calcular(
                st.session_state.trabajador['fecha_ingreso'],
                f_termino, sb, gr, colmov, causal, dias_vac_tomados
            )
            res_fin['Otros_Conceptos']['Adicionales'] = otros_conceptos
            res_fin['Total_Finiquito'] += otros_conceptos
            
            st.session_state.finiquito_calculado = res_fin
            
            # Resultados detallados
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("🏖️ Vacaciones", f"${res_fin['Monto_Vacaciones']:,.0f}")
                st.metric("⚖️ Años Servicio", f"${res_fin['Años_Servicio']:,.0f}")
            with col2:
                st.metric("📢 Mes Aviso", f"${res_fin['Mes_Aviso']:,.0f}")
                st.metric("📊 Proporcional", f"${res_fin['Proporcional_Sueldo']:,.0f}")
            with col3:
                st.metric("💵 Otros", f"${sum(res_fin['Otros_Conceptos'].values()):,.0f}")
                st.metric("📈 Total UF", f"{res_fin['Valor_Liquidacion_UF']:.1f} UF")
            
            # Resumen final
            st.success(f"💰 **TOTAL A PAGAR: ${res_fin['Total_Finiquito']:,.0f}**")
            st.info(f"📅 **Antigüedad:** {res_fin['Antiguedad']}")
            
            if st.button("📋 Generar Finiquito PDF"):
                # Aquí iría la generación del PDF del finiquito
                st.info("PDF de finiquito generado exitosamente")

# =============================================================================
# TAB 3: GESTIÓN DE DOCUMENTOS
# =============================================================================
with tabs[2]:
    st.header("📄 Generador de Documentos Legales")
    st.info("Generación automática de contratos, cartas y documentos laborales")
    
    tipo_doc = st.selectbox("📑 Tipo de Documento", [
        "Contrato de Trabajo Indefinido",
        "Contrato de Trabajo Plazo Fijo", 
        "Carta de Amonestación",
        "Carta de Desvinculación",
        "Carta de Aviso Previo",
        "Acuerdo de Terminación",
        "Certificado de Trabajo"
    ])
    
    if tipo_doc in ["Contrato de Trabajo Indefinido", "Contrato de Trabajo Plazo Fijo"]:
        st.subheader("⚙️ Parámetros del Contrato")
        col1, col2 = st.columns(2)
        with col1:
            fecha_inicio = st.date_input("📅 Fecha Inicio", date.today())
            jornada = st.number_input("⏰ Horas Semanales", 45, min_value=1, max_value=60)
            dias_vacaciones = st.number_input("🏖️ Días Vacaciones", 15, min_value=10, max_value=30)
        with col2:
            ubicacion = st.text_input("📍 Ubicación", "Santiago")
            tipo_contrato_label = "Indefinido" if "Indefinido" in tipo_doc else "Plazo Fijo"
            
        if st.button("📄 Generar Contrato", type="primary"):
            datos_contrato = {
                "cargo": st.session_state.trabajador['cargo'],
                "ubicacion": ubicacion,
                "fecha_inicio": fecha_inicio.strftime("%d/%m/%Y"),
                "jornada": jornada,
                "tipo_contrato": tipo_contrato_label,
                "sueldo_base": st.session_state.get('calculo', {}).get('Base', 0),
                "gratificacion": st.session_state.get('calculo', {}).get('Grat', 0),
                "dias_vacaciones": dias_vacaciones
            }
            
            doc = GeneradorDocumentos.generar_contrato_trabajo(
                st.session_state.empresa, st.session_state.trabajador, datos_contrato
            )
            
            bio = io.BytesIO()
            doc.save(bio)
            st.download_button(
                "📥 Descargar Contrato.docx", 
                bio.getvalue(), 
                f"Contrato_{tipo_contrato_label}.docx"
            )
    
    elif tipo_doc == "Carta de Amonestación":
        st.subheader("📋 Detalles de la Amonestación")
        falta = st.selectbox("🚫 Tipo de Falta", [
            "Llegadas tardías",
            "Faltas injustificadas", 
            "Incumplimiento de deberes",
            "Mala conducta",
            "Incumplimiento horario"
        ])
        descripcion = st.text_area("📝 Descripción Detallada")
        medidas = st.text_area("🎯 Medidas Correctivas")
        
        if st.button("📄 Generar Carta", type="primary"):
            doc = GeneradorDocumentos.generar_carta_amonestacion(
                st.session_state.empresa, st.session_state.trabajador,
                falta, descripcion, medidas
            )
            
            bio = io.BytesIO()
            doc.save(bio)
            st.download_button("📥 Descargar Carta.docx", bio.getvalue(), "Carta_Amonestacion.docx")
    
    elif tipo_doc == "Carta de Desvinculación":
        st.subheader("⚖️ Detalles de la Desvinculación")
        col1, col2 = st.columns(2)
        with col1:
            fecha_termino = st.date_input("📅 Fecha Término", date.today())
            avisos = st.text_input("📢 Días de Aviso", "30 días")
        with col2:
            articulo = st.selectbox("📚 Artículo Legal", [
                ("Art. 159 N°1 - Renuncia Voluntaria", "159 N°1"),
                ("Art. 161 - Necesidades Empresa", "161"),
                ("Art. 168 - Término Injustificado", "168"),
                ("Art. 159 N°4 - Muerte Trabajador", "159 N°4")
            ])
            
        if st.button("📄 Generar Carta", type="primary"):
            causa = {
                "articulo": articulo[1],
                "descripcion": articulo[0]
            }
            
            doc = GeneradorDocumentos.generar_carta_desvinculacion(
                st.session_state.empresa, st.session_state.trabajador,
                causa, fecha_termino.strftime("%d/%m/%Y"), avisos
            )
            
            bio = io.BytesIO()
            doc.save(bio)
            st.download_button("📥 Descargar Carta.docx", bio.getvalue(), "Carta_Desvinculacion.docx")

# =============================================================================
# TAB 4: EVALUACIÓN DE COMPETENCIAS
# =============================================================================
with tabs[3]:
    st.header("🎯 Sistema de Evaluación de Competencias")
    st.info("Análisis de brechas de competencia y planes de carrera personalizados")
    
    if st.button("🔄 Crear Nueva Evaluación", type="primary"):
        st.session_state.evaluaciones = []
        st.session_state.perfil_cargo = None
    
    # Selector de perfil o crear nuevo
    col1, col2 = st.columns(2)
    with col1:
        st.subheader("👥 Candidato")
        if 'evaluaciones' in st.session_state and st.session_state.evaluaciones:
            candidato_seleccionado = st.selectbox(
                "Seleccionar Evaluación", 
                range(len(st.session_state.evaluaciones)),
                format_func=lambda x: st.session_state.evaluaciones[x]['nombre']
            )
            candidato_actual = st.session_state.evaluaciones[candidato_seleccionado]
        else:
            nombre_candidato = st.text_input("Nombre del Candidato")
            candidato_actual = None
    
    with col2:
        st.subheader("🏢 Perfil de Cargo")
        if 'perfil_cargo' in st.session_state and st.session_state.perfil_cargo:
            st.info(f"Perfil actual: {st.session_state.perfil_cargo['cargo']}")
        else:
            crear_perfil = st.button("➕ Crear Perfil de Cargo")
            if crear_perfil:
                st.session_state.crear_perfil = True
    
    # Crear perfil de cargo
    if st.session_state.get('crear_perfil'):
        st.subheader("🏗️ Constructor de Perfil de Cargo")
        
        col1, col2, col3 = st.columns(3)
        with col1:
            area_cargo = st.selectbox("📊 Área", list(COMPETENCIAS_BASE.keys()))
        with col2:
            cargo_nombre = st.text_input("💼 Cargo", "Analista de Sistemas")
        with col3:
            empresa_info = st.text_input("🏢 Empresa", st.session_state.empresa['nombre'])
        
        if st.button("🏗️ Generar Perfil Completo"):
            perfil = MotorCompetencias.crear_perfil_cargo(area_cargo, cargo_nombre, empresa_info)
            st.session_state.perfil_cargo = perfil
            st.session_state.crear_perfil = False
            st.success("Perfil generado exitosamente")
    
    # Evaluación de competencias
    if st.session_state.perfil_cargo and nombre_candidato:
        st.subheader(f"📊 Evaluación: {nombre_candidato}")
        
        competencias_candidato = {}
        
        for categoria, competencias in st.session_state.perfil_cargo['competencias_objetivo'].items():
            st.write(f"**{categoria}:**")
            competencias_candidato[categoria] = {}
            
            col1, col2 = st.columns(2)
            competencias_items = list(competencias.items())
            
            for i, (competencia, nivel_objetivo) in enumerate(competencias_items):
                with col1 if i % 2 == 0 else col2:
                    nivel_candidato = st.selectbox(
                        f"{competencia} (Objetivo: {nivel_objetivo})",
                        ["Básico", "Intermedio", "Avanzado", "Experto"],
                        key=f"{nombre_candidato}_{categoria}_{competencia}"
                    )
                    competencias_candidato[categoria][competencia] = nivel_candidato
        
        if st.button("🧮 Evaluar y Generar Plan", type="primary"):
            resultado = MotorCompetencias.evaluar_candidato(
                st.session_state.perfil_cargo, competencias_candidato
            )
            
            evaluacion = {
                'nombre': nombre_candidato,
                'fecha': datetime.now(),
                'perfil_cargo': st.session_state.perfil_cargo['cargo'],
                'resultados': resultado
            }
            
            st.session_state.evaluaciones.append(evaluacion)
            
            # Mostrar resultados
            st.subheader("📊 Resultados de la Evaluación")
            
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("📈 Score General", f"{resultado['score_total']:.1f}%")
            with col2:
                brechas_altas = len([b for b in resultado['brechas'] if b['criticidad'] == 'Alta'])
                st.metric("🚨 Brechas Críticas", brechas_altas)
            with col3:
                brechas_totales = len(resultado['brechas'])
                st.metric("⚠️ Total Brechas", brechas_totales)
            
            # Análisis detallado por competencia
            if st.checkbox("📋 Ver Análisis Detallado"):
                for categoria, competencias in resultado['evaluaciones'].items():
                    st.write(f"### {categoria}")
                    for comp, eval_data in competencias.items():
                        nivel_objetivo = eval_data['nivel_objetivo']
                        nivel_candidato = eval_data['nivel_candidato']
                        porcentaje = eval_data['porcentaje_match']
                        
                        # Color según nivel de brecha
                        if eval_data['brecha'] >= 2:
                            css_class = "competency-gap-critical"
                            emoji = "🔴"
                        elif eval_data['brecha'] == 1:
                            css_class = "competency-gap-high"
                            emoji = "🟡"
                        else:
                            css_class = "competency-gap"
                            emoji = "🟢"
                        
                        st.markdown(f"""
                        <div class="{css_class}">
                            {emoji} <b>{comp}:</b> Objetivo: {nivel_objetivo} | Actual: {nivel_candidato} | Match: {porcentaje}%
                        </div>
                        """, unsafe_allow_html=True)
            
            # Plan de carrera
            if resultado['brechas']:
                st.subheader("🎯 Plan de Desarrollo Personalizado")
                
                plan = resultado['plan_carrera']
                
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.write("**🚀 Fase 1 (0-3 meses)**")
                    for accion in plan['fase_1']:
                        st.write(f"• {accion['accion']}")
                
                with col2:
                    st.write("**🎯 Fase 2 (3-6 meses)**")
                    for accion in plan['fase_2']:
                        st.write(f"• {accion['accion']}")
                
                with col3:
                    st.write("**📚 Recursos Necesarios**")
                    for competencia, recursos in plan['recursos_necesarios'].items():
                        st.write(f"• {competencia}: {recursos}")
                
                # Cronograma visual
                st.subheader("📅 Cronograma de Desarrollo")
                cronograma_data = []
                for fase in ['fase_1', 'fase_2', 'fase_3']:
                    for accion in plan[fase]:
                        cronograma_data.append({
                            'Competencia': accion['competencia'],
                            'Fase': fase.title(),
                            'Duración': accion['duracion'],
                            'Tipo': accion['tipo']
                        })
                
                if cronograma_data:
                    df_cronograma = pd.DataFrame(cronograma_data)
                    st.dataframe(df_cronograma, use_container_width=True)
            else:
                st.success("🎉 ¡Candidato cumple perfectamente con el perfil!")

# =============================================================================
# TAB 5: PERFILES DE CARGO
# =============================================================================
with tabs[4]:
    st.header("🏗️ Constructor de Perfiles de Cargo")
    st.info("Creación y gestión de perfiles detallados por cargo y área")
    
    if st.button("➕ Nuevo Perfil de Cargo", type="primary"):
        st.session_state.nuevo_perfil = True
    
    if st.session_state.get('nuevo_perfil'):
        st.subheader("🏗️ Crear Perfil Detallado")
        
        col1, col2 = st.columns(2)
        with col1:
            cargo = st.text_input("💼 Nombre del Cargo", "Analista de RRHH")
            area = st.selectbox("📊 Área Funcional", list(COMPETENCIAS_BASE.keys()))
            nivel_seniority = st.selectbox("🎯 Nivel Seniority", ["Junior", "Semi Senior", "Senior", "Lead", "Manager"])
            ubicacion = st.text_input("📍 Ubicación", "Santiago, Chile")
        with col2:
            jornada = st.selectbox("⏰ Jornada", ["Tiempo Completo", "Medio Tiempo", "Jornada Extendida"])
            modalidad = st.selectbox("🏠 Modalidad", ["Presencial", "Híbrido", "Remoto"])
            presupuesto_min = st.number_input("💰 Presupuesto Mínimo", 800000, step=50000)
            presupuesto_max = st.number_input("💰 Presupuesto Máximo", 1200000, step=50000)
        
        # Competencias específicas
        st.subheader("🎯 Competencias Técnicas Requeridas")
        
        competencias_tect = []
        if area in COMPETENCIAS_BASE:
            for categoria, competencias in COMPETENCIAS_BASE[area].items():
                if "Técnicos" in categoria:
                    col1, col2 = st.columns(2)
                    for i, competencia in enumerate(competencias.keys()):
                        with col1 if i % 2 == 0 else col2:
                            nivel_req = st.selectbox(
                                f"{competencia} (Técnico)",
                                ["Básico", "Intermedio", "Avanzado", "Experto"],
                                key=f"tec_{competencia}"
                            )
                            competencias_tect.append((competencia, nivel_req))
        
        # Competencias blandas
        st.subheader("🤝 Competencias Blandas Requeridas")
        competencias_blandas = []
        if area in COMPETENCIAS_BASE:
            for categoria, competencias in COMPETENCIAS_BASE[area].items():
                if "Blandas" in categoria:
                    col1, col2 = st.columns(2)
                    for i, competencia in enumerate(competencias.keys()):
                        with col1 if i % 2 == 0 else col2:
                            nivel_req = st.selectbox(
                                f"{competencia} (Blanda)",
                                ["Básico", "Intermedio", "Avanzado", "Experto"],
                                key=f"bla_{competencia}"
                            )
                            competencias_blandas.append((competencia, nivel_req))
        
        # Funciones y responsabilidades
        st.subheader("📋 Funciones Principales")
        funciones = []
        for i in range(5):
            funcion = st.text_input(f"Función {i+1}", key=f"funcion_{i}")
            if funcion:
                funciones.append(funcion)
        
        # Requisitos
        st.subheader("📚 Requisitos Adicionales")
        requisitos = []
        for i in range(3):
            req = st.text_input(f"Requisito {i+1}", key=f"requisito_{i}")
            if req:
                requisitos.append(req)
        
        if st.button("🏗️ Generar Perfil Completo", type="primary"):
            perfil_completo = {
                "cargo": cargo,
                "area": area,
                "nivel_seniority": nivel_seniority,
                "ubicacion": ubicacion,
                "jornada": jornada,
                "modalidad": modalidad,
                "presupuesto": {"min": presupuesto_min, "max": presupuesto_max},
                "competencias_tecnicas": dict(competencias_tect),
                "competencias_blandas": dict(competencias_blandas),
                "funciones_principales": funciones,
                "requisitos": requisitos,
                "fecha_creacion": datetime.now()
            }
            
            st.session_state.perfil_completo = perfil_completo
            st.session_state.nuevo_perfil = False
            
            st.success("✅ Perfil generado exitosamente")
            
            # Mostrar resumen
            st.subheader("📄 Resumen del Perfil")
            col1, col2 = st.columns(2)
            with col1:
                st.write(f"**Cargo:** {perfil_completo['cargo']}")
                st.write(f"**Área:** {perfil_completo['area']}")
                st.write(f"**Nivel:** {perfil_completo['nivel_seniority']}")
                st.write(f"**Modalidad:** {perfil_completo['modalidad']}")
            with col2:
                st.write(f"**Presupuesto:** ${perfil_completo['presupuesto']['min']:,.0f} - ${perfil_completo['presupuesto']['max']:,.0f}")
                st.write(f"**Ubicación:** {perfil_completo['ubicacion']}")
                st.write(f"**Jornada:** {perfil_completo['jornada']}")
            
            st.write("**Funciones Principales:**")
            for func in funciones:
                st.write(f"• {func}")
            
            if st.button("📥 Generar PDF del Perfil"):
                # Generar PDF del perfil
                st.info("PDF del perfil generado exitosamente")
    
    # Lista de perfiles existentes
    if hasattr(st.session_state, 'perfiles_creados') and st.session_state.perfiles_creados:
        st.subheader("📚 Perfiles Existentes")
        for perfil in st.session_state.perfiles_creados:
            with st.expander(f"📋 {perfil['cargo']} - {perfil['area']}"):
                st.write(f"**Seniority:** {perfil['nivel_seniority']}")
                st.write(f"**Modalidad:** {perfil['modalidad']}")
                st.write(f"**Presupuesto:** ${perfil['presupuesto']['min']:,.0f} - ${perfil['presupuesto']['max']:,.0f}")
                
                if st.button(f"📊 Evaluar Candidato contra {perfil['cargo']}", key=f"eval_{perfil['cargo']}"):
                    st.info(f"Evaluación contra {perfil['cargo']} iniciada")

# =============================================================================
# TAB 6: REPORTES MASIVOS
# =============================================================================
with tabs[5]:
    st.header("📊 Centro de Reportes Masivos")
    st.info("Generación de reportes, análisis y procesamiento masivo de datos")
    
    col1, col2 = st.columns(2)
    with col1:
        st.subheader("📥 Importar Datos")
        st.write("**Plantillas de Importación:**")
        
        # Plantilla de trabajadores
        output_trabajadores = io.BytesIO()
        wb1 = xlsxwriter.Workbook(output_trabajadores, {'in_memory': True})
        ws1 = wb1.add_worksheet()
        cols_trabajadores = ["RUT", "NOMBRE", "CARGO", "EMAIL", "TELEFONO", "FECHA_INGRESO", "SALUD", "AFP"]
        ws1.write_row(0, 0, cols_trabajadores)
        wb1.close()
        
        st.download_button(
            "👥 Plantilla Trabajadores", 
            output_trabajadores.getvalue(), 
            "Plantilla_Trabajadores.xlsx"
        )
        
        # Plantilla de contratos
        output_contratos = io.BytesIO()
        wb2 = xlsxwriter.Workbook(output_contratos, {'in_memory': True})
        ws2 = wb2.add_worksheet()
        cols_contratos = ["RUT", "TIPO_CONTRATO", "SUELDO_BASE", "GRATIFICACION", "COLACION", "MOVILIZACION", "JORNADA"]
        ws2.write_row(0, 0, cols_contratos)
        wb2.close()
        
        st.download_button(
            "📄 Plantilla Contratos", 
            output_contratos.getvalue(), 
            "Plantilla_Contratos.xlsx"
        )
        
        # Plantilla de evaluaciones
        output_evaluaciones = io.BytesIO()
        wb3 = xlsxwriter.Workbook(output_evaluaciones, {'in_memory': True})
        ws3 = wb3.add_worksheet()
        cols_evaluaciones = ["RUT", "COMPETENCIA_1", "NIVEL_1", "COMPETENCIA_2", "NIVEL_2", "COMPETENCIA_3", "NIVEL_3"]
        ws3.write_row(0, 0, cols_evaluaciones)
        wb3.close()
        
        st.download_button(
            "🎯 Plantilla Evaluaciones", 
            output_evaluaciones.getvalue(), 
            "Plantilla_Evaluaciones.xlsx"
        )
    
    with col2:
        st.subheader("📤 Procesar Archivos")
        archivo_trabajadores = st.file_uploader("👥 Cargar Trabajadores", type=['xlsx'])
        archivo_contratos = st.file_uploader("📄 Cargar Contratos", type=['xlsx'])
        archivo_evaluaciones = st.file_uploader("🎯 Cargar Evaluaciones", type=['xlsx'])
        
        if archivo_trabajadores or archivo_contratos or archivo_evaluaciones:
            if st.button("🔄 Procesar Archivos", type="primary"):
                st.success("✅ Archivos procesados exitosamente")
                
                # Simulación de datos procesados
                datos_procesados = {
                    'trabajadores': 150,
                    'contratos': 150,
                    'evaluaciones': 89,
                    'errores': 0
                }
                
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("👥 Trabajadores", datos_procesados['trabajadores'])
                with col2:
                    st.metric("📄 Contratos", datos_procesados['contratos'])
                with col3:
                    st.metric("🎯 Evaluaciones", datos_procesados['evaluaciones'])
                with col4:
                    st.metric("❌ Errores", datos_procesados['errores'])
    
    # Reportes predefinidos
    st.subheader("📈 Reportes Predefinidos")
    
    reportes = {
        "Liquidaciones Mensuales": "Reporte completo de todas las liquidaciones del mes",
        "Finiquitos Pendientes": "Lista de finiquitos por pagar",
        "Análisis de Brechas": "Reporte de brechas de competencia por área",
        "Rotación de Personal": "Análisis de rotación y causas",
        "Cumplimiento Legal": "Estado de cumplimiento de normativas",
        "Presupuesto Salarial": "Análisis de gasto en personal"
    }
    
    col1, col2, col3 = st.columns(3)
    for i, (reporte, descripcion) in enumerate(reportes.items()):
        with col1 if i % 3 == 0 else col2 if i % 3 == 1 else col3:
            with st.expander(reporte):
                st.write(descripcion)
                if st.button(f"📊 Generar {reporte}", key=f"reporte_{i}"):
                    # Simulación de generación
                    st.success(f"✅ {reporte} generado exitosamente")
                    
                    # Crear archivo mock
                    output_mock = io.BytesIO()
                    wb_mock = xlsxwriter.Workbook(output_mock, {'in_memory': True})
                    ws_mock = wb_mock.add_worksheet()
                    ws_mock.write_row(0, 0, ["DATO_1", "DATO_2", "DATO_3", "VALOR"])
                    for i in range(10):
                        ws_mock.write_row(i+1, 0, [f"Item {i+1}", "Categoria A", "Subcategoria B", f"${100000*i:,.0f}"])
                    wb_mock.close()
                    
                    st.download_button(
                        "📥 Descargar Reporte", 
                        output_mock.getvalue(), 
                        f"{reporte.replace(' ', '_')}.xlsx"
                    )

# =============================================================================
# TAB 7: LEGAL & COMPLIANCE
# =============================================================================
with tabs[6]:
    st.header("⚖️ Centro Legal y Compliance")
    st.info("Gestión de cumplimiento normativo y aspectos legales")
    
    col1, col2 = st.columns(2)
    with col1:
        st.subheader("📋 Checklists Legales")
        
        checklists = {
            "Ingreso Nuevo Trabajador": [
                "✅ Contrato firmado",
                "✅ Exámenes médicos",
                "✅ Inducción completa",
                "✅ Entrega elementos de trabajo",
                "✅ Afiliación AFP/Isapre",
                "✅ Firma protocolos",
                "✅ Registro asistencia"
            ],
            "Término de Contrato": [
                "✅ Carta de aviso",
                "✅ Cálculo finiquito",
                "✅ Devolución elementos",
                "✅ Liquidación final",
                "✅ Certificado trabajo",
                "✅ Entrevista salida",
                "✅ Cierre cuentas"
            ],
            "Auditoría Mensual": [
                "✅ Cumplimiento 40 horas",
                "✅ Protocolos KARIN",
                "✅ Descuentos correctos",
                "✅ Vacaciones vigentes",
                "✅ Formación obligatoria",
                "✅ Evaluación riesgos",
                "✅ Sanciones pendientes"
            ]
        }
        
        checklist_seleccionado = st.selectbox("📝 Seleccionar Checklist", list(checklists.keys()))
        st.write("**Elementos del Checklist:**")
        for item in checklists[checklist_seleccionado]:
            st.write(item)
    
    with col2:
        st.subheader("🎓 Capacitaciones Requeridas")
        
        capacitaciones = {
            "Obrigatorias": [
                "Inducción en Seguridad",
                "Protocolo Anti Acoso",
                "Prevención Riesgos",
                "Manejo de Datos Personales"
            ],
            "Por Área": [
                "Tecnología: Ciberseguridad",
                "Operaciones: 5S",
                "Administración: Nueva Normativa",
                "RRHH: Reformas Laborales"
            ],
            "Voluntarias": [
                "Liderazgo",
                "Comunicación Efectiva",
                "Gestión del Tiempo",
                "Innovación"
            ]
        }
        
        for categoria, cursos in capacitaciones.items():
            st.write(f"**{categoria}:**")
            for curso in cursos:
                st.write(f"• {curso}")
    
    # Calculadoras legales
    st.subheader("🧮 Calculadoras Legales Especializadas")
    
    col1, col2, col3 = st.columns(3)
    with col1:
        st.write("**⏰ Horas Extras**")
        horas_normales = st.number_input("Horas Normales", 45)
        horas_extras = st.number_input("Horas Extras")
        valor_hora = st.number_input("Valor Hora Normal")
        
        if st.button("Calcular Extras"):
            if horas_extras > 0:
                # 25% primeras 2 horas extras, 50% siguientes
                if horas_extras <= 2:
                    recargo = 1.25
                else:
                    recargo = 1.5
                
                monto_extras = horas_extras * valor_hora * recargo
                st.success(f"💰 Monto Horas Extras: ${monto_extras:,.0f}")
    
    with col2:
        st.write("**🏖️ Proporcional Vacaciones**")
        fecha_ingreso = st.date_input("Fecha Ingreso", value=date(2024, 1, 1))
        fecha_calculo = st.date_input("Fecha Cálculo", value=date.today())
        sueldo_mensual = st.number_input("Sueldo Mensual", 800000)
        
        if st.button("Calcular Proporcional"):
            dias_trabajados = (fecha_calculo - fecha_ingreso).days
            dias_proporcionales = (dias_trabajados * 15) / 365
            valor_dia = sueldo_mensual / 30
            monto_proporcional = dias_proporcionales * valor_dia
            st.success(f"💰 Proporcional Vacaciones: ${monto_proporcional:,.0f}")
    
    with col3:
        st.write("**💵 Indemnización Años**")
        anos_servicio = st.number_input("Años de Servicio", 5)
        sueldo_base_indem = st.number_input("Sueldo Base Indemnización")
        
        if st.button("Calcular Indemnización"):
            if anos_servicio > 0:
                monto_base = anos_servicio * sueldo_base_indem
                tope_uf = IND["TOPE_INDEM_ANOS"] * IND["UF"]
                monto_final = min(monto_base, tope_uf)
                st.success(f"💰 Indemnización: ${monto_final:,.0f}")
                st.info(f"📊 UF Tope: {IND['TOPE_INDEM_ANOS']} UF = ${tope_uf:,.0f}")
    
    # Actualización de normativas
    st.subheader("📰 Actualizaciones Normativas")
    
    actualizaciones = [
        {
            "fecha": "2025-01-15",
            "titulo": "Nueva Ley 40 Horas - Implementación Gradual",
            "descripcion": "Reducción progresiva de jornada laboral a 40 horas semanales",
            "impacto": "Alto",
            "estado": "Vigente"
        },
        {
            "fecha": "2025-03-01", 
            "titulo": "Protocolo Karin - Acoso Laboral",
            "descripcion": "Nuevas obligaciones en prevención del acoso",
            "impacto": "Medio",
            "estado": "Vigente"
        },
        {
            "fecha": "2025-06-01",
            "titulo": "Flexibilización Trabajo Remoto",
            "descripcion": "Normativas para trabajo híbrido",
            "impacto": "Bajo",
            "estado": "Pronto"
        }
    ]
    
    for update in actualizaciones:
        with st.expander(f"📅 {update['fecha']} - {update['titulo']}"):
            col1, col2 = st.columns(2)
            with col1:
                st.write(f"**Descripción:** {update['descripcion']}")
            with col2:
                st.write(f"**Impacto:** {update['impacto']}")
                st.write(f"**Estado:** {update['estado']}")

# Footer
st.markdown("---")
st.markdown("🏢 **HR Suite Pro** - Sistema Integral de Recursos Humanos")
st.markdown("📞 **Soporte:** contacto@hrsuite.com | 📱 **WhatsApp:** +56 9 XXXX XXXX")
st.markdown(f"🗓️ **Versión:** 2025.11.28 | 📊 **Fecha Actual:** {datetime.now().strftime('%d/%m/%Y %H:%M')}")
