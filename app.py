import streamlit as st
import pandas as pd
import requests
import base64
import os
import io
import zipfile
import random
from datetime import datetime, date
import plotly.graph_objects as go
import plotly.express as px

# =============================================================================
# 1. GESTIÓN DE LIBRERÍAS Y DEPENDENCIAS
# =============================================================================
try:
    import pdfplumber
    from fpdf import FPDF
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import xlsxwriter
    LIBRARIES_OK = True
except ImportError as e:
    LIBRARIES_OK = False
    MISSING_LIB = str(e)

# =============================================================================
# 2. CONFIGURACIÓN DEL SISTEMA
# =============================================================================
st.set_page_config(
    page_title="HR Suite Enterprise V43",
    page_icon="🏢",
    layout="wide",
    initial_sidebar_state="expanded"
)

# =============================================================================
# 3. INICIALIZACIÓN DE ESTADO (MEMORIA DE SESIÓN)
# =============================================================================
def init_state():
    """Inicializa todas las variables de sesión para evitar errores de 'KeyError'."""
    defaults = {
        "empresa": {
            "rut": "", "nombre": "", "giro": "Servicios Profesionales", "direccion": "", 
            "ciudad": "Santiago", "rep_nombre": "", "rep_rut": ""
        },
        "trabajador": {
            "rut": "", "nombre": "", "nacionalidad": "Chilena", 
            "civil": "Soltero", "nacimiento": date(1990, 1, 1), 
            "domicilio": "", "cargo": "", "email": ""
        },
        "calculo_actual": None,
        "finiquito_actual": None,
        "perfil_actual": None,
        "analisis_cv": None,
        "logo_bytes": None
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

init_state()

# =============================================================================
# 4. SISTEMA DE DISEÑO VISUAL (CSS CORPORATIVO)
# =============================================================================
def cargar_estilos():
    nombres = ['fondo.png', 'fondo.jpg', 'fondo_marca.png']
    img = next((n for n in nombres if os.path.exists(n)), None)
    
    css_fondo = ""
    if img:
        try:
            with open(img, "rb") as f: b64 = base64.b64encode(f.read()).decode()
            css_fondo = f"""
            [data-testid="stAppViewContainer"] {{
                background-image: url("data:image/png;base64,{b64}");
                background-size: cover;
                background-position: center;
                background-attachment: fixed;
            }}
            """
        except: pass
    else:
        css_fondo = """[data-testid="stAppViewContainer"] {background: linear-gradient(135deg, #fdfbfb 0%, #ebedee 100%);}"""

    st.markdown(f"""
        <style>
        {css_fondo}
        .block-container {{
            background-color: rgba(255, 255, 255, 0.98); 
            padding: 2.5rem; 
            border-radius: 12px; 
            box-shadow: 0 10px 30px rgba(0,0,0,0.1);
        }}
        h1, h2, h3, h4 {{color: #004a99 !important; font-family: 'Segoe UI', sans-serif; font-weight: 800;}}
        p, label, li {{color: #003366 !important; font-weight: 500;}}
        
        /* Botones Acción */
        .stButton>button {{
            background: linear-gradient(90deg, #004a99 0%, #003366 100%);
            color: white !important;
            font-weight: bold;
            border-radius: 6px;
            height: 3rem;
            width: 100%;
            border: none;
            box-shadow: 0 4px 6px rgba(0,0,0,0.1);
            text-transform: uppercase;
            letter-spacing: 1px;
            transition: 0.3s;
        }}
        .stButton>button:hover {{transform: translateY(-2px); box-shadow: 0 6px 12px rgba(0,0,0,0.2);}}
        
        /* Liquidación Visual */
        .liq-paper {{
            border: 1px solid #ccc; padding: 30px; background: #fff;
            font-family: 'Courier New', monospace; margin-top: 20px;
            box-shadow: 5px 5px 15px rgba(0,0,0,0.05);
        }}
        .liq-header {{text-align: center; border-bottom: 2px solid #000; padding-bottom: 10px; margin-bottom: 20px; font-weight: bold;}}
        .liq-row {{display: flex; justify-content: space-between; border-bottom: 1px dotted #ccc; padding: 5px 0;}}
        .liq-total {{
            background: #e3f2fd; padding: 15px; font-weight: bold; font-size: 1.2em;
            border: 2px solid #004a99; margin-top: 20px; color: #004a99; text-align: right;
        }}
        .miles-feedback {{color: #28a745; font-weight: bold; font-size: 0.85em; margin-top: -10px;}}
        </style>
        """, unsafe_allow_html=True)

cargar_estilos()

# =============================================================================
# 5. UTILIDADES Y DATA OFICIAL
# =============================================================================
def fmt(valor):
    if pd.isna(valor) or valor == "": return "$0"
    try: return "${:,.0f}".format(float(valor)).replace(",", ".")
    except: return str(valor)

def mostrar_feedback(valor):
    if valor > 0: st.markdown(f'<p class="miles-feedback">✅ Ingresaste: {fmt(valor)}</p>', unsafe_allow_html=True)

def obtener_indicadores_oficiales():
    """Valores Oficiales Previred Noviembre 2025[cite: 8, 9, 13]."""
    return {
        "UF": 39643.59,
        "UTM": 69542.0,
        "SUELDO_MIN": 529000,
        "TOPE_AFP_UF": 87.8,
        "TOPE_AFC_UF": 131.9,
        "TOPE_INDEM_UF": 90
    }

IND = obtener_indicadores_oficiales()

# =============================================================================
# 6. MOTOR DE CÁLCULO (ISAPRE TARGET & REVERSO)
# =============================================================================
def calcular_sueldo_liquido_a_bruto(liquido_objetivo, colacion, movilizacion, tipo_contrato, afp_nombre, salud_sistema, plan_uf_isapre):
    
    no_imponibles = colacion + movilizacion
    liquido_tributable_meta = liquido_objetivo - no_imponibles
    
    if liquido_tributable_meta < IND["SUELDO_MIN"] * 0.4: return None

    # Topes y Constantes
    TOPE_GRAT_MENSUAL = (4.75 * IND["SUELDO_MIN"]) / 12
    TOPE_IMP_PESOS = IND["TOPE_AFP_UF"] * IND["UF"]
    TOPE_AFC_PESOS = IND["TOPE_AFC_UF"] * IND["UF"]
    
    # Tasas AFP Nov 2025 [cite: 16]
    TASAS_AFP_DICT = {"Capital": 11.44, "Cuprum": 11.44, "Habitat": 11.27, "PlanVital": 11.16, "Provida": 11.45, "Modelo": 10.58, "Uno": 10.46, "SIN AFP": 0.0}
    
    es_empresarial = (tipo_contrato == "Sueldo Empresarial")
    
    if es_empresarial:
        tasa_afp = 0.0
    else:
        tasa_total = TASAS_AFP_DICT.get(afp_nombre, 11.44)
        tasa_afp = 0.0 if afp_nombre == "SIN AFP" else (tasa_total / 100)

    # Tasas AFC [cite: 23]
    tasa_afc_trab = 0.006 if (tipo_contrato == "Indefinido" and not es_empresarial) else 0.0
    tasa_afc_emp = 0.024
    if not es_empresarial:
        tasa_afc_emp = 0.024 if tipo_contrato == "Indefinido" else (0.03 if tipo_contrato == "Plazo Fijo" else 0.0)

    # Costos Empresa
    tasa_sis = 0.0149 # [cite: 17]
    tasa_mutual = 0.0093 # Base

    # Búsqueda Binaria
    min_bruto = 100000
    max_bruto = liquido_tributable_meta * 2.5
    
    for _ in range(150):
        base_test = (min_bruto + max_bruto) / 2
        gratificacion = min(base_test * 0.25, TOPE_GRAT_MENSUAL)
        if es_empresarial: gratificacion = 0
        
        tot_imponible = base_test + gratificacion
        base_prev = min(tot_imponible, TOPE_IMP_PESOS)
        base_afc = min(tot_imponible, TOPE_AFC_PESOS)
        
        monto_afp = int(base_prev * tasa_afp)
        monto_afc = int(base_afc * tasa_afc_trab)
        legal_7 = int(base_prev * 0.07) # Base Legal 7%
        
        # Impuesto
        base_tributable = max(0, tot_imponible - monto_afp - legal_7 - monto_afc)
        impuesto = 0
        if base_tributable > 13.5 * IND["UTM"]: impuesto = int(base_tributable * 0.04) # Simplificado T1
        if base_tributable > 30 * IND["UTM"]: impuesto = int((base_tributable * 0.08) - (1.74 * IND["UTM"])) # Simplificado T2
        impuesto = max(0, impuesto)
        
        liquido_calc = tot_imponible - monto_afp - legal_7 - monto_afc - impuesto
        
        if abs(liquido_calc - liquido_tributable_meta) < 50:
            # AJUSTE REAL ISAPRE
            salud_descuento = legal_7
            adicional_isapre = 0
            warning_msg = None
            
            if salud_sistema == "Isapre (UF)":
                valor_plan_pesos = int(plan_uf_isapre * IND["UF"])
                if valor_plan_pesos > legal_7:
                    salud_descuento = valor_plan_pesos
                    adicional_isapre = valor_plan_pesos - legal_7
                    warning_msg = f"⚠️ Plan Isapre excede el 7% legal. El líquido final disminuye en {fmt(adicional_isapre)}."
            
            liq_final_real = tot_imponible - monto_afp - salud_descuento - monto_afc - impuesto + no_imponibles
            
            # Costos Empresa
            aporte_sis = int(base_prev * tasa_sis)
            aporte_afc = int(base_afc * tasa_afc_emp)
            aporte_mut = int(base_prev * tasa_mutual)
            total_aportes = aporte_sis + aporte_afc + aporte_mut
            
            return {
                "Sueldo Base": int(base_test),
                "Gratificación": int(gratificacion),
                "Total Imponible": int(tot_imponible),
                "No Imponibles": int(no_imponibles),
                "LÍQUIDO_FINAL": int(liq_final_real),
                "AFP": monto_afp,
                "Salud": salud_descuento,
                "Adicional_Isapre": adicional_isapre,
                "AFC": monto_afc,
                "Impuesto": impuesto,
                "Total Descuentos": monto_afp + salud_descuento + monto_afc + impuesto,
                "Aportes Empresa": total_aportes,
                "COSTO TOTAL": int(tot_imponible + no_imponibles + total_aportes),
                "Warning": warning_msg
            }
            break
        elif liquido_calc < liquido_tributable_meta: min_bruto = base_test
        else: max_bruto = base_test
            
    return None

def calcular_finiquito_legal(f_ini, f_fin, sueldo_base, causal, vac_pend):
    """Calcula indemnizaciones con topes legales[cite: 12]."""
    dias_trabajados = (f_fin - f_ini).days
    anos_servicio = dias_trabajados / 365.25
    anos_pago = int(anos_servicio)
    if (anos_servicio - anos_pago) * 12 >= 6: anos_pago += 1
    if anos_pago > 11: anos_pago = 11 
    
    tope_indem = IND["TOPE_INDEM_UF"] * IND["UF"]
    base_calc = min(sueldo_base, tope_indem)
    
    monto_anos = int(base_calc * anos_pago) if causal == "Necesidades de la Empresa" else 0
    monto_aviso = int(base_calc) if causal in ["Necesidades de la Empresa", "Desahucio"] else 0
    
    valor_dia = sueldo_base / 30
    monto_feriado = int(vac_pend * 1.25 * valor_dia)
    
    return {
        "Años Servicio": monto_anos, "Aviso Previo": monto_aviso, "Vacaciones": monto_feriado,
        "TOTAL": monto_anos + monto_aviso + monto_feriado, "Antigüedad": f"{anos_pago} años"
    }

# =============================================================================
# 7. GENERADORES DOCUMENTALES (WORD/PDF/EXCEL)
# =============================================================================

def generar_contrato_word_full(fin, emp, trab, cond):
    if not LIBRARIES_OK: return None
    doc = Document()
    style = doc.styles['Normal']; style.font.name = 'Arial'; style.font.size = Pt(11)
    
    t = doc.add_heading('CONTRATO DE TRABAJO INDEFINIDO', 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")
    
    fecha_hoy = datetime.now().strftime("%d de %B de %Y")
    
    intro = f"""En {emp.get('ciudad','Santiago')}, a {fecha_hoy}, entre la empresa "{emp.get('nombre','').upper()}", RUT {emp.get('rut','')}, giro {emp.get('giro','Servicios')}, representada legalmente por don/ña {emp.get('rep_nombre','').upper()}, RUT {emp.get('rep_rut','')}, ambos domiciliados en {emp.get('direccion','')}, en adelante el "EMPLEADOR"; y don/ña {trab.get('nombre','').upper()}, RUT {trab.get('rut','')}, de nacionalidad {trab.get('nacionalidad','')}, nacido el {str(trab.get('nacimiento',''))}, domiciliado en {trab.get('domicilio','')}, en adelante el "TRABAJADOR", se ha convenido:"""
    
    p = doc.add_paragraph(intro); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # CLÁUSULAS ROBUSTAS
    clausulas = [
        ("PRIMERO (Cargo y Funciones):", f"El Trabajador se desempeñará como {cond['cargo'].upper()}, realizando las siguientes funciones principales: {cond['funciones']}."),
        ("SEGUNDO (Lugar de Trabajo):", f"Los servicios se prestarán en las dependencias de la empresa ubicadas en {emp.get('direccion','')}, sin perjuicio de los desplazamientos requeridos."),
        ("TERCERO (Jornada):", "El trabajador cumplirá una jornada ordinaria de 44 horas semanales, distribuidas de lunes a viernes (Sujeto a reducción gradual Ley 40 Horas)."),
        ("CUARTO (Remuneración):", f"El Empleador pagará:\na) Sueldo Base Mensual: {fmt(fin['Sueldo Base'])}\nb) Gratificación Legal: {fmt(fin['Gratificación'])} (Tope 4.75 IMM)\nc) Asignación Colación: {fmt(fin['No Imponibles']//2)}\nd) Asignación Movilización: {fmt(fin['No Imponibles']//2)}"),
        ("QUINTO (Confidencialidad):", "El Trabajador guardará estricta reserva de la información confidencial de la Empresa."),
        ("SEXTO (Propiedad Intelectual):", "Toda creación intelectual desarrollada en el marco del contrato será propiedad exclusiva del Empleador."),
        ("SÉPTIMO (Ley Karin):", "La empresa cuenta con un Protocolo de Prevención del Acoso Sexual, Laboral y Violencia, el cual es conocido por el trabajador."),
        ("OCTAVO (Vigencia):", f"El presente contrato es de carácter {cond['tipo']} y comenzará a regir a partir del {str(cond['inicio'])}.")
    ]
    
    for tit, txt in clausulas:
        p = doc.add_paragraph(); r = p.add_run(tit); r.bold = True; p.add_run(f" {txt}")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; doc.add_paragraph("")
        
    doc.add_paragraph("\n\n__________________________\nFIRMA EMPLEADOR\n\n__________________________\nFIRMA TRABAJADOR").alignment = WD_ALIGN_PARAGRAPH.CENTER
    bio = io.BytesIO(); doc.save(bio); return bio

def generar_pdf_liquidacion(res, emp, trab):
    if not LIBRARIES_OK: return None
    pdf = FPDF()
    pdf.add_page()
    
    # Logo
    if st.session_state.logo_bytes:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
            tmp.write(st.session_state.logo_bytes); tmp_path = tmp.name
        try: pdf.image(tmp_path, 10, 8, 30)
        except: pass

    pdf.set_font("Arial", 'B', 14); pdf.cell(0, 15, "LIQUIDACION DE SUELDO", 0, 1, 'C'); pdf.ln(10)
    pdf.set_font("Arial", '', 10)
    pdf.cell(0, 6, f"Empresa: {emp.get('nombre','')} | RUT: {emp.get('rut','')}", 0, 1)
    pdf.cell(0, 6, f"Trabajador: {trab.get('nombre','')} | RUT: {trab.get('rut','')}", 0, 1)
    pdf.cell(0, 6, f"Fecha: {datetime.now().strftime('%d/%m/%Y')}", 0, 1); pdf.ln(5)
    
    pdf.set_fill_color(240, 240, 240); pdf.set_font("Arial", 'B', 10)
    pdf.cell(95, 8, "HABERES", 1, 0, 'C', 1); pdf.cell(95, 8, "DESCUENTOS", 1, 1, 'C', 1)
    pdf.set_font("Arial", '', 10)
    
    h = [("Sueldo Base", res['Sueldo Base']), ("Gratificacion", res['Gratificación']), ("No Imponibles", res['No Imponibles'])]
    d = [("AFP", res['AFP']), ("Salud", res['Salud']), ("Adicional Isapre", res.get('Adicional_Isapre', 0)), ("Seguro Cesantia", res['AFC']), ("Impuesto", res['Impuesto'])]
    
    for i in range(max(len(h), len(d))):
        ht, hv = h[i] if i<len(h) else ("", "")
        dt, dv = d[i] if i<len(d) else ("", "")
        pdf.cell(60, 6, ht, 'L'); pdf.cell(35, 6, fmt(hv) if hv!="" else "", 'R')
        pdf.cell(60, 6, dt, 'L'); pdf.cell(35, 6, fmt(dv) if dv!="" else "", 'R', 1)
        
    pdf.ln(5); pdf.set_font("Arial", 'B', 12)
    pdf.cell(130, 10, "A PAGAR:", 1, 0, 'R'); pdf.cell(60, 10, fmt(res['LÍQUIDO_FINAL']), 1, 1, 'C')
    return pdf.output(dest='S').encode('latin-1')

def generar_plantilla_excel():
    df = pd.DataFrame(columns=["TIPO_DOCUMENTO", "NOMBRE", "RUT", "CARGO", "SUELDO_BASE", "FECHA_INICIO"])
    df.loc[0] = ["Contrato", "Juan Perez", "12.345.678-9", "Analista", 800000, "2025-01-01"]
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer: df.to_excel(writer, index=False)
    buffer.seek(0)
    return buffer

def procesar_lote_masivo(df, empresa_data):
    zip_buffer = io.BytesIO()
    log = []
    with zipfile.ZipFile(zip_buffer, "w") as zf:
        for idx, row in df.iterrows():
            try:
                nombre = str(row.get('NOMBRE', f'Trab_{idx}'))
                # Simulación de datos para generación masiva
                dummy_fin = {"Sueldo Base": row.get('SUELDO_BASE', 0), "Gratificación": 0, "No Imponibles": 0}
                dummy_trab = {"nombre": nombre, "rut": str(row.get('RUT','')), "nacionalidad": "Chilena", "nacimiento": "", "domicilio": ""}
                dummy_cond = {"cargo": str(row.get('CARGO','')), "funciones": "N/A", "tipo": "Indefinido", "inicio": row.get('FECHA_INICIO', date.today())}
                
                doc = generar_contrato_word_full(dummy_fin, empresa_data, dummy_trab, dummy_cond)
                zf.writestr(f"Contrato_{nombre}.docx", doc.getvalue())
            except Exception as e: log.append(f"Fila {idx}: {e}")
    zip_buffer.seek(0)
    return zip_buffer, log

# =============================================================================
# 8. PERFIL DE CARGO (ESTRUCTURA WORD)
# =============================================================================
def generar_perfil_word_style(cargo, rubro):
    if not cargo: return None
    return {
        "titulo": cargo.title(),
        "objetivo": f"Gestionar y controlar los procesos críticos de {cargo} en el sector {rubro}.",
        "dependencia": "Gerencia General / Gerencia de Área",
        "nivel_resp": "Jefatura / Supervisión Senior",
        "funciones": [
            "Coordinación y supervisión de equipos de trabajo.",
            "Control presupuestario y gestión de recursos.",
            "Reportabilidad de estados de avance y KPIs.",
            "Aseguramiento de la calidad y normativa vigente."
        ],
        "requisitos": [
            "Título Profesional Universitario.",
            f"Experiencia mínima de 4 años en {rubro}.",
            "Manejo de ERP y Office Avanzado.",
            "Inglés Técnico."
        ],
        "competencias": ["Liderazgo", "Autonomía", "Orientación al Resultado", "Trabajo bajo presión"],
        "condiciones": ["Jornada: Art. 22 / 44 Horas.", "Lugar: Oficina Central / Terreno.", "Disponibilidad viajes."]
    }

# =============================================================================
# 9. INTERFAZ GRÁFICA (DASHBOARD)
# =============================================================================

# --- SIDEBAR ---
with st.sidebar:
    st.image("https://www.previred.com/wp-content/uploads/2021/01/logo-previred.png", width=120)
    
    st.markdown("### 🏢 Configuración Empresa")
    up_logo = st.file_uploader("Logo (Para PDF)", type=["png", "jpg"])
    if up_logo: st.session_state.logo_bytes = up_logo.read()
    
    with st.expander("Datos Empresa (Fijos)", expanded=True):
        st.session_state.empresa['rut'] = st.text_input("RUT Empresa", st.session_state.empresa['rut'])
        st.session_state.empresa['nombre'] = st.text_input("Razón Social", st.session_state.empresa['nombre'])
        st.session_state.empresa['rep_nombre'] = st.text_input("Rep. Legal", st.session_state.empresa['rep_nombre'])
        st.session_state.empresa['rep_rut'] = st.text_input("RUT Rep.", st.session_state.empresa['rep_rut'])
        st.session_state.empresa['direccion'] = st.text_input("Dirección", st.session_state.empresa['direccion'])
        st.session_state.empresa['ciudad'] = st.text_input("Ciudad", st.session_state.empresa['ciudad'])
        st.session_state.empresa['giro'] = st.text_input("Giro", st.session_state.empresa['giro'])

    with st.expander("👤 Datos Trabajador (Opcional)", expanded=False):
        st.session_state.trabajador['nombre'] = st.text_input("Nombre", st.session_state.trabajador['nombre'])
        st.session_state.trabajador['rut'] = st.text_input("RUT Trab.", st.session_state.trabajador['rut'])
        st.session_state.trabajador['direccion'] = st.text_input("Domicilio", st.session_state.trabajador['direccion'])
        st.session_state.trabajador['nacimiento'] = st.date_input("Nacimiento", value=date(1990,1,1))

    st.divider()
    st.metric("UF Hoy", fmt(IND['UF']))

st.title("HR Suite Enterprise V43")
st.markdown("**Sistema Integral de Gestión de Personas**")

tabs = st.tabs(["💰 Calculadora", "📂 Carga Masiva", "📋 Perfil Cargo", "📜 Legal Hub", "📊 Indicadores"])

# --- TAB 1: CALCULADORA ---
with tabs[0]:
    c1, c2 = st.columns(2)
    with c1:
        liq = st.number_input("Líquido Objetivo", 1000000, step=50000); mostrar_feedback(liq)
        col = st.number_input("Colación", 50000); mov = st.number_input("Movilización", 50000)
    with c2:
        con = st.selectbox("Contrato", ["Indefinido", "Plazo Fijo", "Empresarial"])
        afp = st.selectbox("AFP", ["Capital", "Habitat", "Modelo", "Uno", "SIN AFP"])
        sal = st.selectbox("Salud", ["Fonasa (7%)", "Isapre (UF)"])
        plan = st.number_input("Plan UF", 0.0) if sal == "Isapre (UF)" else 0.0

    if st.button("CALCULAR NÓMINA"):
        res = calcular_sueldo_liquido_a_bruto(liq, col, mov, con, afp, sal, plan)
        if res:
            st.session_state.calculo_actual = res
            if res.get('Warning'): st.warning(res['Warning'])
            
            k1, k2, k3 = st.columns(3)
            k1.metric("Sueldo Base", fmt(res['Sueldo Base']))
            k2.metric("Líquido", fmt(res['LÍQUIDO_FINAL']))
            k3.metric("Costo Empresa", fmt(res['COSTO TOTAL']))
            
            st.markdown(f"""
            <div class="liq-paper">
                <div class="liq-header">LIQUIDACIÓN SIMULADA</div>
                <div class="liq-row"><span>Sueldo Base:</span><span>{fmt(res['Sueldo Base'])}</span></div>
                <div class="liq-row"><span>Gratificación:</span><span>{fmt(res['Gratificación'])}</span></div>
                <div class="liq-row"><span>No Imponibles:</span><span>{fmt(res['No Imponibles'])}</span></div>
                <hr>
                <div class="liq-row"><span>AFP:</span><span style="color:red">-{fmt(res['AFP'])}</span></div>
                <div class="liq-row"><span>Salud:</span><span style="color:red">-{fmt(res['Salud'])}</span></div>
                {f'<div class="liq-row"><span>Adicional Isapre:</span><span style="color:red">-{fmt(res["Adicional_Isapre"])}</span></div>' if res['Adicional_Isapre'] > 0 else ''}
                <div class="liq-row"><span>Impuesto:</span><span style="color:red">-{fmt(res['Impuesto'])}</span></div>
                <div class="liq-total">A PAGAR: {fmt(res['LÍQUIDO_FINAL'])}</div>
            </div>
            """, unsafe_allow_html=True)
            
            if LIBRARIES_OK:
                pdf = generar_pdf_liquidacion(res, st.session_state.empresa, st.session_state.trabajador)
                st.download_button("⬇️ Descargar PDF", pdf, "liquidacion.pdf", "application/pdf")

# --- TAB 2: CARGA MASIVA ---
with tabs[1]:
    st.header("Generación Masiva")
    if LIBRARIES_OK:
        plantilla = generar_plantilla_excel()
        st.download_button("1. Descargar Plantilla Excel", plantilla, "plantilla.xlsx")
        
        up_file = st.file_uploader("2. Subir Excel", type="xlsx")
        if up_file and st.button("PROCESAR Y GENERAR ZIP"):
            if st.session_state.empresa['rut']:
                df = pd.read_excel(up_file)
                zip_file, errs = procesar_lote_masivo(df, st.session_state.empresa)
                st.success(f"Procesado. Errores: {len(errs)}")
                st.download_button("⬇️ Descargar ZIP", zip_file, "docs.zip", "application/zip")
            else: st.error("Faltan datos empresa.")

# --- TAB 3: PERFIL ---
with tabs[2]:
    cargo = st.text_input("Cargo", "Analista")
    rubro = st.selectbox("Rubro", ["Minería", "TI", "Retail"])
    if cargo:
        perf = generar_perfil_robusto(cargo, rubro)
        st.info(perf['objetivo'])
        c1, c2 = st.columns(2)
        c1.write("**Funciones:**"); c1.write("\n".join([f"- {x}" for x in perf['funciones']]))
        c2.write("**Requisitos:**"); c2.write("\n".join([f"- {x}" for x in perf['requisitos']]))

# --- TAB 4: LEGAL HUB ---
with tabs[3]:
    modo = st.radio("Documento", ["Contrato Individual", "Finiquito"])
    
    if modo == "Contrato Individual":
        if st.session_state.calculo_actual:
            if st.button("Generar Contrato"):
                if st.session_state.empresa['rut']:
                    cond = {"cargo": cargo, "funciones": "Las del cargo", "tipo": "Indefinido", "inicio": date.today()}
                    doc = generar_contrato_word_full(st.session_state.calculo_actual, st.session_state.empresa, st.session_state.trabajador, cond)
                    st.download_button("Descargar Word", doc.getvalue(), "contrato.docx")
                else: st.error("Faltan datos empresa")
        else: st.info("Calcule sueldo en Pestaña 1 primero")
        
    elif modo == "Finiquito":
        c1, c2 = st.columns(2)
        fi = c1.date_input("Inicio", date(2020,1,1)); ft = c2.date_input("Fin", date.today())
        base = st.number_input("Base Finiquito", 1000000)
        vac = st.number_input("Vacaciones Pendientes", 0.0)
        causal = st.selectbox("Causal", ["Renuncia", "Necesidades"])
        
        if st.button("Calcular Finiquito"):
            res = calcular_finiquito_legal(fi, ft, base, causal, vac)
            st.write(res)

# --- TAB 5: INDICADORES ---
with tabs[4]:
    st.header("Indicadores Previred (Nov 2025)")
    st.info(f"UF: {fmt(IND['UF'])} | UTM: {fmt(IND['UTM'])}")
    st.table(pd.DataFrame({"Concepto": ["Sueldo Mínimo", "Tope AFP (UF)", "Tope AFC (UF)"], "Valor": [fmt(IND['SUELDO_MIN']), IND['TOPE_AFP_UF'], IND['TOPE_AFC_UF']]}))
